"""
Resume parsing module for extracting text from various file formats
"""
import os
import PyPDF2
from docx import Document
import logging
from typing import Optional

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeParser:
    """Handles parsing of resume files in various formats"""
    
    def __init__(self):
        """Initialize resume parser"""
        pass
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """Extract text from resume file based on file extension"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            logger.info(f"Successfully extracted text from PDF: {os.path.basename(file_path)}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            # Try alternative method
            return self._extract_from_pdf_alternative(file_path)
    
    def _extract_from_pdf_alternative(self, file_path: str) -> str:
        """Alternative PDF extraction method"""
        try:
            import fitz  # PyMuPDF - install if needed
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text.strip()
        except ImportError:
            logger.warning("PyMuPDF not available for alternative PDF extraction")
            return ""
        except Exception as e:
            logger.error(f"Alternative PDF extraction failed: {str(e)}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {os.path.basename(file_path)}")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Successfully extracted text from TXT: {os.path.basename(file_path)}")
            return text.strip()
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.info(f"Successfully extracted text from TXT using {encoding}: {os.path.basename(file_path)}")
                    return text.strip()
                except:
                    continue
            
            logger.error(f"Could not decode text file {file_path}")
            return ""
            
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {str(e)}")
            return ""
    
    def is_valid_resume(self, text: str) -> bool:
        """Check if extracted text appears to be a valid resume"""
        if not text or len(text.strip()) < 50:
            return False
        
        # Common resume indicators
        resume_indicators = [
            'experience', 'education', 'skills', 'work', 'employment',
            'university', 'college', 'degree', 'certification', 'project',
            'responsibility', 'achievement', 'qualification'
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in resume_indicators if indicator in text_lower)
        
        # Consider valid if at least 3 indicators are present
        return indicator_count >= 3
